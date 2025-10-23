import os
import shutil
import subprocess
from datetime import datetime
from pathlib import Path
from typing import Tuple
from docxtpl import DocxTemplate
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.counterparty import Counterparty

CONTRACTS_DIR = "../../../contracts"
DOCX_TEMPLATE_PATH = "../../../Договор_поставки_ТиО_2020_с_клиентом.docx"
RFQ_TEMPLATE_PATH = "../../../contracts/Запрос-КП-от-поставщика.docx"

def ensure_contracts_dir():
    Path(CONTRACTS_DIR).mkdir(exist_ok=True)

def _get_month_name_in_russian(month_number: int) -> str:
    months = {
        1: "января", 2: "февраля", 3: "марта", 4: "апреля", 5: "мая", 6: "июня",
        7: "июля", 8: "августа", 9: "сентября", 10: "октября", 11: "ноября", 12: "декабря"
    }
    return months.get(month_number, "")

def _prepare_context(counterparty: Counterparty) -> dict:
    current_date = datetime.now()
    month_name = _get_month_name_in_russian(current_date.month)

    return {
        "contract_number": counterparty.id,
        "date": f"«{current_date.day:02}» {month_name} {current_date.year} г.",
        "short_name": counterparty.short_name or "",
        "inn": counterparty.inn or "",
        "kpp": counterparty.kpp or "",
        "ogrn": counterparty.ogrn or "",
        "legal_address": counterparty.legal_address or "",
        "bank_name": counterparty.bank_name or "",
        "bank_bik": counterparty.bank_bik or "",
        "bank_account": counterparty.bank_account or "",
        "bank_corr": counterparty.bank_corr or "",
        "director": getattr(counterparty, "director", "") or "",
        "phone": getattr(counterparty, "phone", "") or "",
        "email": getattr(counterparty, "email", "") or "",
    }


def _clear_old_files(counterparty_id: int):
    for filename in os.listdir(CONTRACTS_DIR):
        if filename.startswith(f"contract_{counterparty_id}_"):
            os.remove(os.path.join(CONTRACTS_DIR, filename))


def _convert_to_pdf(docx_path: str, pdf_path: str) -> bool:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice:
        try:
            subprocess.run([soffice, "--headless", "--convert-to", "pdf", "--outdir", os.path.dirname(pdf_path), docx_path], check=True)
            return os.path.exists(pdf_path)
        except Exception:
            pass
    try:
        from docx2pdf import convert as docx2pdf_convert
        docx2pdf_convert(docx_path, pdf_path)
        return os.path.exists(pdf_path)
    except Exception:
        return False


async def generate_contract_for_counterparty(counterparty_id: int, db: AsyncSession) -> Tuple[str, str]:
    counterparty = await db.get(Counterparty, counterparty_id)
    if not counterparty:
        raise ValueError(f"Counterparty with ID {counterparty_id} not found")

    ensure_contracts_dir()
    _clear_old_files(counterparty_id)

    context = _prepare_context(counterparty)
    
    doc = DocxTemplate(DOCX_TEMPLATE_PATH)
    doc.render(context)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    docx_filename = f"contract_{counterparty_id}_{timestamp}.docx"
    pdf_filename = f"contract_{counterparty_id}_{timestamp}.pdf"

    docx_path = os.path.join(CONTRACTS_DIR, docx_filename)
    pdf_path = os.path.join(CONTRACTS_DIR, pdf_filename)

    doc.save(docx_path)
    ok = _convert_to_pdf(docx_path, pdf_path)
    return docx_path, pdf_path if ok else ""

async def generate_supplier_rfq_for_counterparty(counterparty_id: int, db: AsyncSession) -> Tuple[str, str]:
    counterparty = await db.get(Counterparty, counterparty_id)
    if not counterparty:
        raise ValueError(f"Counterparty with ID {counterparty_id} not found")

    ensure_contracts_dir()
    _clear_old_files(counterparty_id)

    context = _prepare_context(counterparty)

    doc = Document(RFQ_TEMPLATE_PATH)

    def replace_in_paragraphs(paragraphs):
        for p in paragraphs:
            for r in p.runs:
                if r.font.highlight_color == WD_COLOR_INDEX.GREEN:
                    key = r.text.strip().strip('{}').strip().lower()
                    if key in context:
                        r.text = str(context[key])
                        r.font.highlight_color = None

    replace_in_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_paragraphs(cell.paragraphs)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    docx_filename = f"rfq_{counterparty_id}_{timestamp}.docx"
    pdf_filename = f"rfq_{counterparty_id}_{timestamp}.pdf"

    docx_path = os.path.join(CONTRACTS_DIR, docx_filename)
    pdf_path = os.path.join(CONTRACTS_DIR, pdf_filename)

    doc.save(docx_path)
    ok = _convert_to_pdf(docx_path, pdf_path)
    return docx_path, pdf_path if ok else ""