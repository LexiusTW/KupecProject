import io
import os
import shutil
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

from fastapi import APIRouter, Depends, File, HTTPException, UploadFile, status
from fastapi.responses import FileResponse
from pydantic import BaseModel
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.organization import Organization
from app.api.deps import get_db, get_current_user
from app.models.user import User
from app.schemas.excel import FileListResponse

router = APIRouter(prefix="/docs", tags=["docs"])


STATIC_DIR = Path(__file__).resolve().parents[4] / "static"
LOGO_DIR = STATIC_DIR / "logos"
TEMPLATES_DIR = Path(__file__).resolve().parents[4] / "contracts" / "templates"
OUT_DIR = Path(__file__).resolve().parents[4] / "excel_outgoing"

ALLOWED_IMG = {"image/png", "image/jpeg", "image/webp", "image/svg+xml"}


@router.post("/logo")
async def upload_logo(
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user),
):
    if file.content_type not in ALLOWED_IMG:
        raise HTTPException(status_code=415, detail="Поддерживаются только PNG/JPEG/WEBP/SVG")
    
    # Загружаем организацию, чтобы получить ИНН и сохранить logo_url
    organization = await db.get(Organization, user.organization_id)
    if not organization:
        raise HTTPException(status_code=404, detail="Организация пользователя не найдена")
    if not organization.inn:
        raise HTTPException(status_code=400, detail="У организации не заполнен ИНН")

    LOGO_DIR.mkdir(parents=True, exist_ok=True)
    ext = {
        "image/png": ".png",
        "image/jpeg": ".jpg",
        "image/webp": ".webp",
        "image/svg+xml": ".svg",
    }.get(file.content_type, ".png")

    filename = f"logo_{organization.inn}{ext}"
    dest_path = LOGO_DIR / filename

    with dest_path.open("wb") as f:
        shutil.copyfileobj(file.file, f)

    organization.logo_url = f"/static/logos/{filename}"
    await db.commit()

    return {"logo_url": organization.logo_url}


class DocItem(BaseModel):
    name: str
    qty: float
    unit: str
    price: float | None = None
    comment: str | None = None


class GenerateDocPayload(BaseModel):
    template_name: str
    org_code: str | None = None
    emp_code: str | None = None
    category_code: str | None = None
    seq_num: int | None = None
    year: int | None = None

    metal_items: List[DocItem] | None = None
    other_items: List[DocItem] | None = None
    variables: Dict[str, Any] | None = None


def _two_digit_year(year: Optional[int]) -> str:
    if year is None:
        year = datetime.now().year % 100
    else:
        year = year % 100
    return f"{year:02d}"


def _build_number(org: Optional[str], emp: Optional[str], year2: str, seq: int, cat: Optional[str]) -> str:
    org = (org or "ORG").upper()
    emp = (emp or "EMP").upper()
    cat = (cat or "XX").upper()
    return f"{org}-{emp}-{year2}-{seq}-{cat}"


def _render_docx(template_path: Path, ctx: Dict[str, Any], out_docx: Path) -> None:
    try:
        from docx import Document  # python-docx
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"docx модуль недоступен: {e}")

    doc = Document(str(template_path))
    def replace_in_text(text: str) -> str:
        out = text
        for k, v in ctx.items():
            out = out.replace(f"{{{{{k}}}}}", str(v))
        return out

    for p in doc.paragraphs:
        if p.text and "{{" in p.text:
            for run in p.runs:
                run.text = replace_in_text(run.text)

    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                if cell.text and "{{" in cell.text:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.text = replace_in_text(run.text)

    def build_table(placeholder: str, items: Optional[List[DocItem]]):
        if not items:
            for p in doc.paragraphs:
                if placeholder in p.text:
                    p.text = p.text.replace(placeholder, "")
            return
        inserted = False
        for i, p in enumerate(doc.paragraphs):
            if placeholder in p.text:
                p.text = p.text.replace(placeholder, "")
                table = doc.add_table(rows=1, cols=5)
                hdr = table.rows[0].cells
                hdr[0].text = "Наименование"
                hdr[1].text = "Кол-во"
                hdr[2].text = "Ед."
                hdr[3].text = "Цена"
                hdr[4].text = "Комментарий"
                for it in items:
                    row = table.add_row().cells
                    row[0].text = it.name
                    row[1].text = str(it.qty)
                    row[2].text = it.unit
                    row[3].text = "" if it.price is None else str(it.price)
                    row[4].text = it.comment or ""
                inserted = True
                break
        if not inserted:
            pass

    build_table("{{TABLE_METAL}}", ctx.get("metal_items"))
    build_table("{{TABLE_OTHER}}", ctx.get("other_items"))

    doc.save(str(out_docx))


def _convert_to_pdf(src_docx: Path, out_pdf: Path) -> bool:
    try:
        from docx2pdf import convert
    except Exception:
        return False
    try:
        convert(str(src_docx), str(out_pdf))
        return out_pdf.exists()
    except Exception:
        return False


@router.post("/generate")
async def generate_document(
    payload: GenerateDocPayload,
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user),
):
    year2 = _two_digit_year(payload.year)
    seq = payload.seq_num or 1
    number = _build_number(payload.org_code, payload.emp_code, year2, seq, payload.category_code)

    ctx: Dict[str, Any] = {}
    if payload.variables:
        ctx.update(payload.variables)

    organization = await db.get(Organization, user.organization_id)
    if not organization:
        raise HTTPException(status_code=404, detail="Организация пользователя не найдена")

    ctx.update({
        "org_login": user.login,
        "org_email": user.email,
        "org_inn": organization.inn or "",
        "org_director": organization.director_name or "",
        "org_legal_address": organization.legal_address or "",
        "org_ogrn": organization.ogrn or "",
        "org_kpp": organization.kpp or "",
        "org_okpo": organization.okpo or "",
        "org_okato_oktmo": organization.okato_oktmo or "",
        "logo_url": organization.logo_url or "",
        "doc_number": number,
        "doc_date": datetime.now().strftime("%d.%m.%Y"),
    })

    metal_items = payload.metal_items or []
    other_items = payload.other_items or []
    ctx["metal_items"] = [it.model_dump() for it in metal_items] if metal_items else None
    ctx["other_items"] = [it.model_dump() for it in other_items] if other_items else None

    template_path = (TEMPLATES_DIR / payload.template_name).resolve()
    if not template_path.exists():
        raise HTTPException(status_code=404, detail="Шаблон не найден")

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_docx = OUT_DIR / f"generated_{stamp}.docx"
    out_pdf = OUT_DIR / f"generated_{stamp}.pdf"

    _render_docx(template_path, ctx, out_docx)

    if _convert_to_pdf(out_docx, out_pdf):
        return FileResponse(str(out_pdf), media_type="application/pdf", filename=out_pdf.name)
    else:
        return FileResponse(str(out_docx), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=out_docx.name)


@router.get("/generated", response_model=FileListResponse)
async def list_generated_documents(
    user: User = Depends(get_current_user),
):
    """Возвращает список сгенерированных документов из папки excel_outgoing."""
    if not OUT_DIR.exists() or not OUT_DIR.is_dir():
        return FileListResponse(directory=str(OUT_DIR), files=[], total_count=0)

    try:
        # Получаем список файлов, отсортированный по времени изменения (новые вверху)
        files = sorted(
            (f for f in OUT_DIR.iterdir() if f.is_file()),
            key=lambda f: f.stat().st_mtime,
            reverse=True
        )
        filenames = [f.name for f in files]
        return FileListResponse(directory=str(OUT_DIR.name), files=filenames, total_count=len(filenames))
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Ошибка при чтении директории: {e}"
        )


@router.get("/download/{filename}")
async def download_generated_document(
    filename: str,
    user: User = Depends(get_current_user),
):
    """Скачивание сгенерированного документа по имени файла."""
    file_path = (OUT_DIR / filename).resolve()

    # Проверка безопасности, чтобы нельзя было выйти за пределы OUT_DIR
    if not str(file_path).startswith(str(OUT_DIR.resolve())):
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Доступ запрещен")

    if not file_path.exists() or not file_path.is_file():
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Файл не найден")

    return FileResponse(path=file_path, filename=filename)
