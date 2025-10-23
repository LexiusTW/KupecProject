from pathlib import Path

from fastapi import APIRouter, Depends, HTTPException, status, Request
from pydantic import BaseModel
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.api.deps import get_db, get_current_user
from app.models.user import User
from app.services.contracts import generate_contract_for_counterparty, generate_supplier_rfq_for_counterparty, ensure_contracts_dir
import os
import glob
from app.schemas.contract import ContractGenerationResponse, ContractDataSchema
from app.models.contract import Contract

router = APIRouter()

@router.get("/check/{counterparty_id}", response_model=ContractGenerationResponse)
async def check_contract_endpoint(
    counterparty_id: int,
    request: Request,
    user: User = Depends(get_current_user),
    
):
    if user.role == "seller":
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Sellers cannot check contracts")

    contracts_dir = os.path.join(os.getcwd(), "contracts")
    pattern = os.path.join(contracts_dir, f"contract_{counterparty_id}_*.pdf")
    
    matching_files = glob.glob(pattern)

    if matching_files:
        file_path = matching_files[0]
        file_name = os.path.basename(file_path)
        contract_url = request.url_for("contracts", path=file_name)
        
        return ContractGenerationResponse(
            counterparty_id=counterparty_id,
            file_path=str(contract_url),
            message="Contract PDF found"
        )
    else:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Contract PDF not found for this counterparty.")

@router.post("/generate/{counterparty_id}", response_model=ContractGenerationResponse)
async def generate_contract_endpoint(
    counterparty_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user),
):
    ensure_contracts_dir()

    if user.role == "seller":
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Sellers cannot generate contracts")

    try:
        docx_path, pdf_path = await generate_contract_for_counterparty(counterparty_id, db)
        
        if pdf_path:
            file_name = os.path.basename(pdf_path)
            contract_url = request.url_for("contracts", path=file_name)
        elif docx_path:
            file_name = os.path.basename(docx_path)
            contract_url = request.url_for("contracts", path=file_name)
        else:
            contract_url = None

        return ContractGenerationResponse(
            counterparty_id=counterparty_id,
            file_path=str(contract_url) if contract_url else None,
            message="Contract generated successfully"
        )
    except ValueError as e:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"Error generating contract: {str(e)}")


@router.patch("/edit/{contract_id}")
async def edit_contract(
        contract_id: int,
        data: ContractDataSchema,
        db: AsyncSession = Depends(get_db)
):
    from docx import Document

    result = await db.execute(select(Contract).where(Contract.id == contract_id))
    contract = result.scalar_one_or_none()

    if not contract:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Contract not found")

    path = Path(__file__).resolve().parents[4] / "contracts"
    doc_path = path / "doc.docx" # путь до шаблона

    document = Document(str(doc_path))

    current_data = ContractDataSchema(**(contract.data or {}))
    updated_data = current_data.model_copy(update=data.model_dump(exclude_unset=True))
    contract.data = updated_data.model_dump()

    await db.commit()
    await db.refresh(contract)


    def replace_in_text(text: str) -> str:
        out = text
        for k, v in contract.data.items():
            out = out.replace(f"{{{{{k}}}}}", str(v) if v is not None else "Не указано")  # Заменяем {{key}}
        return out

    # --- Заменяем в параграфах ---
    for p in document.paragraphs:
        full_text = "".join(run.text for run in p.runs)
        if "{{" in full_text:
            new_text = replace_in_text(full_text)
            for i, run in enumerate(p.runs):
                run.text = "" if i > 0 else new_text  # очищаем остальные runs

    # --- Заменяем в таблицах ---
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    full_text = "".join(run.text for run in p.runs)
                    if "{{" in full_text:
                        new_text = replace_in_text(full_text)
                        for i, run in enumerate(p.runs):
                            run.text = "" if i > 0 else new_text

    # --- Сохраняем итоговый документ ---
    document.save(str(path / f"contract_{contract.id}.docx"))

    return {"status": "success"}
@router.post("/generate-rfq/{counterparty_id}", response_model=ContractGenerationResponse)
async def generate_rfq_endpoint(
    counterparty_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user),
):
    ensure_contracts_dir()

    if user.role == "seller":
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Sellers cannot generate RFQ")

    try:
        docx_path, pdf_path = await generate_supplier_rfq_for_counterparty(counterparty_id, db)
        if pdf_path:
            file_name = os.path.basename(pdf_path)
            contract_url = request.url_for("contracts", path=file_name)
        elif docx_path:
            file_name = os.path.basename(docx_path)
            contract_url = request.url_for("contracts", path=file_name)
        else:
            contract_url = None
        return ContractGenerationResponse(
            counterparty_id=counterparty_id,
            file_path=str(contract_url) if contract_url else None,
            message="RFQ generated successfully"
        )
    except ValueError as e:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"Error generating RFQ: {str(e)}")
