import os
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

CONTRACTS_DIR = "/Users/elr1c180/Desktop/KupecProject/contracts"
RFQ_TEMPLATE_PATH = "/Users/elr1c180/Desktop/KupecProject/contracts/Запрос-КП-от-поставщика.docx"


def ensure_dir():
    Path(CONTRACTS_DIR).mkdir(exist_ok=True)


def build_context():
    now = datetime.now()
    return {
        "contract_number": "TEST-9999",
        "date": now.strftime("%d.%m.%Y"),
        "short_name": "Тест ООО",
        "inn": "7701234567",
        "kpp": "770101001",
        "ogrn": "1234567890123",
        "legal_address": "123456, г. Москва, ул. Тестовая, д.1",
        "bank_name": "ПАО Сбербанк",
        "bank_bik": "044525225",
        "bank_account": "40702810900000000001",
        "bank_corr": "30101810400000000225",
        "director": "Иванов И.И.",
        "phone": "+7 (999) 111-22-33",
        "email": "test@example.com",
    }


def replace_in_paragraphs(paragraphs, context):
    for p in paragraphs:
        for r in p.runs:
            if r.font.highlight_color == WD_COLOR_INDEX.GREEN:
                key = r.text.strip().strip('{}').strip().lower()
                if key in context:
                    r.text = str(context[key])
                    r.font.highlight_color = None


def main():
    ensure_dir()
    if not os.path.exists(RFQ_TEMPLATE_PATH):
        raise FileNotFoundError(RFQ_TEMPLATE_PATH)
    ctx = build_context()
    doc = Document(RFQ_TEMPLATE_PATH)
    replace_in_paragraphs(doc.paragraphs, ctx)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_paragraphs(cell.paragraphs, ctx)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_path = os.path.join(CONTRACTS_DIR, f"rfq_highlight_test_{ts}.docx")
    doc.save(out_path)
    print(out_path)


if __name__ == "__main__":
    main()

